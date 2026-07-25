"""`/report/draft` — a Markdown five-pillar report skeleton.

Auto-fills measured-layer statistics (each with its n-size) and cited-layer entries
(each with its citation), and leaves explicit [ANALYST INPUT] markers wherever human
judgment is required. Trend sub-sections are generated from the project's configured
trend terms PLUS emergent themes found in analysis — never a hard-coded trend list.
"""
from __future__ import annotations

from datetime import datetime, timezone
from typing import Any, Dict, List

import analytics
import storage
from version import __version__

MARKER = "[ANALYST INPUT]"


def _n(label: str, n: int) -> str:
    flag = "  ⚠️ _low-confidence (n<100)_" if n < analytics.LOW_CONFIDENCE_THRESHOLD else ""
    return f"(n={n}){flag}"


def draft_report(project_id: int) -> str:
    project = storage.get_project(project_id)
    if not project:
        raise ValueError(f"Project {project_id} not found")
    cfg = project["config"]
    market = cfg.get("market", {})
    product = cfg.get("product", {})

    dash = analytics.dashboard(project_id)
    by_channel = analytics.sentiment_by_channel(project_id)
    drivers = analytics.top_purchase_drivers(project_id)
    bvc = analytics.brand_vs_competitor_sentiment(project_id)
    trends = analytics.trend_volume_over_time(project_id)
    verbatims = analytics.top_verbatims_per_theme(project_id, per_theme=3)
    cited = storage.list_market_intel(project_id, entry_type="cited")

    lines: List[str] = []
    w = lines.append

    w(f"# Market & Product Intelligence — {product.get('brand', '(brand)')}")
    w(f"_Market: {market.get('country', '')} · Languages: {', '.join(market.get('languages', []))} · "
      f"Generated {datetime.now(timezone.utc).date().isoformat()} · MarketLens v{__version__}_")
    w("")
    w(f"> Draft skeleton. Measured stats carry their n-size; cited stats carry their source. "
      f"Fill every {MARKER} with analyst judgment before delivery.")
    w("")

    # -------------------------------------------------------------- Pillar 1
    w("## 1. Market Overview")
    w("")
    if cited:
        w("**Cited market data:**")
        for e in cited:
            src = f"{e.get('source_name', '?')}, {e.get('publication_date', 'n.d.')}"
            w(f"- **{e.get('metric') or e.get('category')}**: {e.get('value')} "
              f"— _{src}_ ([source]({e.get('source_url', '')})); confidence: {e.get('confidence')}")
    else:
        w(f"- {MARKER}: No cited market data entered yet. Add market size / CAGR / share via the "
          f"Market Intelligence (Cited) workspace.")
    w("")
    w(f"- Total signals collected across channels: **{dash['total_items']}**; analyzed: "
      f"**{dash['total_analyzed']}**.")
    w(f"- {MARKER}: Synthesize the competitive landscape and market structure.")
    w("")

    # -------------------------------------------------------------- Pillar 2
    w("## 2. Consumer Intelligence")
    w("")
    w(f"- Overall net sentiment across analyzed signals: **{dash['overall_net_score']}** "
      f"{_n('overall', dash['total_analyzed'])}.")
    w("")
    w("**Sentiment by channel:**")
    for ch in by_channel:
        w(f"- {ch['channel']}: net **{ch['net_score']}** {_n(ch['channel'], ch['n'])} "
          f"· languages: {ch['language_breakdown']}")
    w("")
    w("**Brand vs. competitor sentiment:**")
    for row in bvc:
        w(f"- {row['brand_focus']}: net **{row['net_score']}** {_n(row['brand_focus'], row['n'])}")
    w("")
    w("**Top purchase drivers** " + _n("drivers", drivers["n"]) + ":")
    for d in drivers["drivers"][:10]:
        w(f"- {d['driver']} ({d['count']})")
    w("")
    w(f"- {MARKER}: Interpret what drives choice and how the brand is perceived vs. rivals.")
    w("")

    # -------------------------------------------------------------- Pillar 3
    w("## 3. Key Trends")
    w("")
    w("_Sub-sections below are generated from configured trend terms + emergent themes in the data._")
    w("")
    for s in trends["series"]:
        theme = s["trend_category"]
        w(f"### Trend: {theme} {_n(theme, s['n'])}")
        months = ", ".join(f"{m}:{c}" for m, c in list(s["by_month"].items())[-6:])
        if months:
            w(f"- Volume over recent months: {months}")
        # Attach up to 3 representative verbatims for this theme.
        theme_v = next((t for t in verbatims["themes"] if t["theme"] == theme), None)
        if theme_v:
            for v in theme_v["verbatims"]:
                if v.get("summary_en"):
                    w(f"  - _\"{v['summary_en']}\"_ ({v['sentiment']}, {v['source']})")
        w(f"- {MARKER}: Is this trend rising, and what does it mean for the brand?")
        w("")

    # -------------------------------------------------------------- Pillar 4
    w("## 4. Product Innovation")
    w("")
    w(f"- {MARKER}: Translate purchase drivers and complaints into product/packaging opportunities.")
    w("- Signals to mine: negative-sentiment verbatims and unmet-need themes above.")
    w("")

    # -------------------------------------------------------------- Pillar 5
    w("## 5. Partnership Opportunities")
    w("")
    manual_ads = storage.list_market_intel(project_id, entry_type="manual_ad")
    if manual_ads:
        w("**Observed advertisers/creatives (Manual Intelligence):**")
        for e in manual_ads[:15]:
            w(f"- {e.get('source_name')}: {e.get('value')} — {e.get('extra', {}).get('creative_theme', '')}")
    else:
        w(f"- {MARKER}: No manual ad observations recorded yet.")
    w(f"- {MARKER}: Identify retail, channel, and co-marketing partnership angles.")
    w("")

    w("---")
    w("### Methodology & honesty notes")
    w(f"- Model-generated tags: ~85–95% human agreement — spot-check before quoting.")
    w(f"- Digital sources skew urban/online/literate-in-covered-languages — not the whole market.")
    w(f"- Tier-3 platforms (Instagram, Facebook, LinkedIn, WhatsApp, TikTok organic, app-only "
      f"delivery) are NOT covered; see the export Methodology tab.")
    return "\n".join(lines)


# --------------------------------------------------------------------------- #
# File outputs (Markdown + Word)
# --------------------------------------------------------------------------- #
def _out_path(project_id: int, ext: str) -> str:
    from pathlib import Path

    from settings import settings

    settings.ensure_dirs()
    project = storage.get_project(project_id)
    safe = "".join(c for c in (project["name"] if project else "project") if c.isalnum() or c in "-_") or "project"
    ts = datetime.now(timezone.utc).strftime("%Y%m%dT%H%M%SZ")
    return str(Path(settings.exports_dir) / f"MarketLens_Report_{safe}_{ts}.{ext}")


def save_markdown(project_id: int, out_path: str = None) -> str:
    md = draft_report(project_id)
    path = out_path or _out_path(project_id, "md")
    with open(path, "w", encoding="utf-8") as f:
        f.write(md)
    storage.audit("report.export.md", "report draft downloaded (markdown)", project_id=project_id)
    return path


def save_docx(project_id: int, out_path: str = None) -> str:
    """Render the Markdown report into a styled .docx (Word) file.

    A lightweight line-based Markdown converter handles the report's structure
    (headings, bullets, block-quotes, horizontal rules, bold). python-docx is
    imported lazily so the rest of the tool never requires it.
    """
    from docx import Document
    from docx.shared import Pt, RGBColor

    md = draft_report(project_id)
    path = out_path or _out_path(project_id, "docx")

    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    def _add_runs(paragraph, text):
        # Minimal **bold** handling; everything else is plain text.
        parts = text.split("**")
        for i, part in enumerate(parts):
            if not part:
                continue
            run = paragraph.add_run(part)
            run.bold = (i % 2 == 1)

    for raw in md.splitlines():
        line = raw.rstrip()
        if not line.strip():
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("# "):
            h = doc.add_heading(line[2:], level=1)
        elif line.strip() == "---":
            doc.add_paragraph().add_run("_" * 40)
        elif line.lstrip().startswith(("- ", "* ")):
            indent = len(line) - len(line.lstrip())
            p = doc.add_paragraph(style="List Bullet")
            if indent >= 2:
                p.paragraph_format.left_indent = Pt(18)
            _add_runs(p, line.lstrip()[2:])
        elif line.startswith(">"):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(18)
            r = p.add_run(line.lstrip("> ").strip())
            r.italic = True
        else:
            p = doc.add_paragraph()
            _add_runs(p, line)

    doc.save(path)
    storage.audit("report.export.docx", "report draft downloaded (Word)", project_id=project_id)
    return path
